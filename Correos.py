import os
import docx
import shutil
import codecs
import smtplib
import pandas as pd
from os import walk
from getpass import getpass
from os.path import basename
from email.mime.text import MIMEText
from email.message import EmailMessage
from cryptography.fernet import Fernet
from email.mime.multipart import MIMEMultipart
from email.utils import COMMASPACE, formatdate
from email.mime.application import MIMEApplication

# ------------------------------------------VARIABLES------------------------------------------

lines = '-'*60
autosave = ""
loginText = ""
optionLine = "Select an Option: "
goBackOption = "0_ Go back"
htmlSelected = ''
option = 0
outOption1 = 0
outOption2 = 0
outOption3 = 0
outOption4 = 0
config = False
empty = False
login = False
data_filesRemaining = []
attachmentSelected = 0
attachmentsSelected = []
numbers = ("0", "1", "2", "3", "4", "5", "6", "7", "8", "9")
abc = ["A", "B", "C", "D", "E", "F", "G", "H",
       "I", "J", "K", "L", "M", "N", "O", "P",
       "Q", "R", "S", "T", "U", "V", "X", "Y", "Z"]

# ------------------------------------------FUNCTIONS------------------------------------------


def clear(): return os.system("cls")


def writeOptionData(*args):
    textFile = os.path.abspath(os.path.join(
        os.path.dirname(
            __file__), 'Options.txt'))

    with open(textFile, "w") as f:
        for i in args:
            f.write(i)


def readText(index):
    filePath = os.path.abspath(os.path.join(
        os.path.dirname(
            __file__), 'Options.txt'))

    with open(filePath, "r") as f:
        file = f.readlines()

    return file[index]


def writeInText(LineIndex, Text):
    file = os.path.abspath(os.path.join(
        os.path.dirname(
            __file__), 'Options.txt'))

    with open(file) as f:
        lines = f.readlines()
    lines[LineIndex] = Text
    with open(file, "w") as f:
        f.writelines(lines)


def getFilesInFolder(path):
    f = []
    names = []
    for (dirpath, dirnames, filenames) in walk(path):
        f.extend(filenames)
        names = filenames
        break
    return names


def getFoldersInFolder(path):
    f = []
    for (dirpath, dirnames, filenames) in walk(path):
        f.extend(dirnames)
        break
    return dirnames


def getDocsInFolder(path):
    count = []
    files = getFilesInFolder(path)
    for file in files:
        file = str(file)
        name, extension = os.path.splitext(path+file)
        if extension == ".docx":
            count.append(file)
    return count

def getHTMLSInFolder(path):
    count = []
    files = getFilesInFolder(path)
    for file in files:
        file = str(file)
        name, extension = os.path.splitext(path+file)
        if extension == ".html":
            count.append(file)
    return count

def checkExcelInFolder(path):
    count = []
    files = getFilesInFolder(path)
    for file in files:
        file = str(file)
        name, extension = os.path.splitext(path+file)
        if extension == ".xlsx":
            count.append(file)
    return count


def createTextFile(name):
    txt = open(name+'.txt', 'a')
    txt.close()


def createOptionsText(path):
    txt = open(path, 'r')
    txt.close()


def createFolderIN(path, name):
    try:
        os.chdir(path)
        os.makedirs(name)
    except:
        pass


def createFolderON(name):
    try:
        path = os.path.dirname(
            __file__)
        os.chdir(path)
        os.makedirs(name)
    except:
        pass


def removeTextFiles(*args):
    for i in args:
        os.remove(i+'.txt')


def getDataFrom():

    excelsFolder = os.path.abspath(
        os.path.join(os.path.dirname(
            __file__), 'Excel Repository'))

    excelsInFolder = checkExcelInFolder(excelsFolder)

    excelSelected = []

    clear()
    print(lines)
    print("We found", len(excelsInFolder),
          "excels in the folder, select the ones you want to get data from")
    print(lines)
    input("Press enter to continue... ")
    clear()

    for excel in excelsInFolder:
        excelsFolder = os.path.abspath(
            os.path.join(os.path.dirname(
                __file__), 'Excel Repository'))
        error = False
        clear()
        print(lines)
        print("Get data from", excel+'?')
        print(lines)
        print("1_ Yes")
        print("2_ No")
        print(lines)
        try:
            option = int(input(optionLine))
        except:
            clear()
            print(lines)
            print("Get data from", excel+'?')
            print(lines)
            print("1_ Yes")
            print("2_ No")
            print(lines)
            option = int(input("Select a correct option: "))
        if option > 2 or option < 1:
            error = True

        while error == True:
            clear()
            print(lines)
            print("Invalid Option, please select a number between 1 and 2")
            print(lines)
            print("1_ Yes")
            print("2_ No")
            print(lines)
            option = int(input(optionLine))
            if numerodesc >= 0:
                error = False

        if option == 1:
            excelSelected.append(excel)

        elif option == 2:
            pass

    data_dir = os.path.abspath(os.path.join(os.path.dirname(
        __file__), 'Data'))

    for ex in excelSelected:
        excelFile = os.path.abspath(os.path.join(
            os.path.dirname(
                __file__), 'Excel Repository', ex))

        excelFile = r"{}".format(excelFile)

        ex = ex[:len(ex)-5]

        createFolderIN(data_dir, ex)

        excelFolder = os.path.abspath(os.path.join(os.path.dirname(
            __file__), "Data", ex))

        cont = 0
        text_file = ""
        txt = 'New.txt'
        file_list = []

        for car in abc:
            cont = 0
            df = pd.read_excel(excelFile, sheet_name=0, usecols=car)
            df.to_csv(txt, header=True, index=False)
            fo = open(txt, 'r')
            content = fo.readlines()

            if (os.stat(txt).st_size > 2) == False:
                fo.close()
                empty == True
                os.remove(txt)

            for i in content:
                var = str(i)
                if cont == 0:
                    l = len(var)
                    rem_last = var[:l-1]
                    createTextFile(rem_last.replace(" ", ""))
                    text_file = rem_last.replace(" ", "")+".txt"
                    file_list.append(str(text_file))
                    cont = 1
                else:
                    add = open(text_file, 'a')
                    add.write(var)
                    add.close()
            fo.close()

        file_list = set(file_list)
        file_list.remove(".txt")
        for file in file_list:
            try:
                shutil.move(file, excelFolder)
            except shutil.Error:
                length = len(file)
                file_woTXT = file[:length-4]
                print(file_woTXT, "--- Already Uploaded")
                os.remove(file)
        os.remove(".txt")


def enter():
    succesConection = True
    clear()
    print(lines)
    print("Login with your gmail")
    print(lines)
    mail = input("Mail: ")
    password = getpass()
    print(lines)
    try:
        # Inicio de sesion
        conection.login(user=mail, password=password)
        print("Valid Credentials")

    except:
        input("Wrong Credentials... ")
        succesConection = False

    while succesConection == False:
        enter()
    return mail, password


def removeExtensions(array):
    array_list = []
    for i in array:
        length = len(i)
        i = i[:length-4]
        array_list.append(i)
    return array_list


def printArray(array):
    array_numbers = 0
    for i, j in enumerate(array):
        print(str(i+1)+'_ '+str(j))
        array_numbers += 1
    return array_numbers


def dataBase():
    mainPath = os.path.abspath(os.path.join(os.path.dirname(
        __file__), 'Data'))
    return getFoldersInFolder(mainPath)


def sumTextToList(dataSelectedFolder, *args):
    dataList = []
    array = []
    for i in args:
        array.append(i)
    for textFiles in array:
        path = os.path.abspath(os.path.join(
            os.path.dirname(
                __file__), 'Data', dataSelectedFolder, textFiles))
        text = open(path, "r").readlines()
        for lines in textFiles:
            dataList.append(lines)
        text.close()
    return print(dataList)


def sendMassiveMails(selectedDataBase, textFiles, wordFile, mails, mail, attachmentsSelected):
    textFilesWE = removeExtensions(textFiles)
    body = ""
    allFiles = []
    mailWT = []
    number = 0
    mails = path_data = os.path.abspath(os.path.join(
        os.path.dirname(
            __file__), 'Data', selectedDataBase, mails))
    for file in textFiles:
        file = os.path.abspath(os.path.join(
            os.path.dirname(
                __file__), 'Data', selectedDataBase, file))
        semiList = []
        file = open(file, "r").readlines()
        for word in file:
            semiList.append(word.strip())
        allFiles.append(semiList)
    mails = open(mails, "r").readlines()
    for i in mails:
        mailWT.append(i.strip())
    for mn, ma in enumerate(mailWT):
        first, second = readDocs(wordFile)
        body = subject = msg = ''
        for i, j in enumerate(textFilesWE):
            j.replace(".", "")
            for m, p in enumerate(first):
                if p == '['+j+']':
                    first[m] = allFiles[i][mn]
                if p == '['+j+'].':
                    first[m] = allFiles[i][mn]+'.'
                if p == '['+j+'],':
                    first[m] = allFiles[i][mn]+','
            for f, g in enumerate(second):
                if g == '['+j+']':
                    second[f] = allFiles[i][mn]
                if g == '['+j+'].':
                    second[f] = allFiles[i][mn]+'.'
                if g == '['+j+'],':
                    second[f] = allFiles[i][mn]+','
        subject = ' '.join(first)

        for i, word in enumerate(second):
            if word != "FLAG":
                body += word + " "
            else:
                body += "<br>"

        with codecs.open(htmlFolder, "r", "utf-8") as f:
            codeHTML = f.read()

        msg = MIMEMultipart()
        msg['Subject'] = subject
        msg['From'] = mail
        msg['To'] = ma

        #msg.attach(MIMEText(body.encode('utf-8'), _charset='utf-8'))
        #msg = f'Subject: {subject.decode("utf-8")}\n\n{body.decode("utf-8")}'
        for files in attachmentsSelected or []:
            attachmentFile = os.path.abspath(os.path.join(os.path.dirname(__file__), 'Attachments', files))
            with open(attachmentFile, "rb") as fil:
                part = MIMEApplication(
                    fil.read(),
                    Name=basename(attachmentFile)
                )
            part['Content-Disposition'] = 'attachment; filename="%s"' % basename(files)
            msg.attach(part)
        part2 = MIMEText(codeHTML.format(body=body), 'html')
        msg.attach(part2)
        conection.sendmail(mail, ma, msg.as_string())
    clear()
    print(lines)
    print("Mails Sended")
    print(lines)
    input("Press Enter to exit... ")


def readDocs(docFile):
    paragraphsList = []
    paragraphsSplited = []
    totalParagraph = []
    path = os.path.abspath(
        os.path.join(os.path.dirname(
            __file__), 'Mails', docFile))
    doc = docx.Document(path)
    firstParagraph = doc.paragraphs[0].text.split()
    paragraphs = doc.paragraphs
    for i,j in enumerate(paragraphs[1:]):
        paragraphsList.append(doc.paragraphs[i+1].text +"\n")
    for elements in paragraphsList:
        paragraphsSplited.append(elements.split())
    for i in paragraphsSplited:
        i.append("FLAG")
        for j in i:
            totalParagraph.append(j)
    return firstParagraph, totalParagraph

def checkFolders(*args):
    path = os.path.abspath(os.path.join(os.path.dirname(__file__)))
    folders = getFoldersInFolder(path)
    for folder in args:
        if folder not in folders:
            createFolderON(folder)


# ------------------------------------------START/SETUP------------------------------------------


try:
    # Establecer conection con el servidor de SMTP Gmail
    conection = smtplib.SMTP_SSL(host='smtp.gmail.com', port=465)
except:
    print("Check your wifi conection")
os.chdir(os.path.dirname(
    __file__))

createTextFile("Options")



if os.path.getsize("Options.txt") == 0:
    criptoKey = Fernet.generate_key()
    writeOptionData("firstTimeOption1 = True\n",
                    "AutoSave = [Disabled]\n", "login=False\n", "Mail: \n", "Password: \n", "HTML \n", criptoKey.decode("utf-8"))

if readText(5) != "HTML \n":
    htmlSelected = readText(5)[:-1]

if readText(3) != "Mail: \n" and readText(4) != "Password: \n":
    crypter = Fernet(readText(-1).encode())
    mail_encrypted = readText(3).encode()
    password_encrypted = readText(4).encode()
    mail = crypter.decrypt(mail_encrypted).decode()
    password = crypter.decrypt(password_encrypted).decode()
    conection.login(user=mail, password=password)
    login = True

autosave = readText(1)[:len(readText(1))-1]

if readText(2) == "login=False\n":
    loginText = "Login"
else:
    loginText = "Logout"

if readText(1) == "AutoSave = [Disabled]\n":
    savePassword = False
elif readText(1) == "AutoSave = [Enabled]\n":
    savePassword = True

checkFolders("Excel Repository","Attachments","Data","Mails","Mails/html templates")

#createFolderON("Excel Repository")
#createFolderON("Attachments")
#createFolderON("Data")
#createFolderON("Mails")

htmlFolder = os.path.abspath(os.path.join(os.path.dirname(__file__), 'Mails','html templates', htmlSelected))
htmlsFolder = os.path.abspath(os.path.join(os.path.dirname(__file__), 'Mails','html templates'))


# ------------------------------------------FUNCTIONS------------------------------------------
while option != 5:
    outOption4 = 0
    outOption3 = 0
    outOption2 = 0
    outOption1 = 0
    clear()
    menu = f'{lines}\n' \
        '-----------------Welcome to the Mail Script-----------------\n'\
        f'{lines}\n' \
        '1_ Upload an Excel file with Data\n'\
        '2_ Send Mails\n'\
        '3_ Options\n'\
        f'{lines}'\

    loginMenu = f'{lines}\n' \
        '------------------------Login Configs-----------------------\n'\
        f'{lines}\n' \
        f'1_ ' + autosave + '\n'\
        '2_ Select HTML Template\n'\
        f'{goBackOption}\n'\
        f'{lines}'\

    print(menu)
    option = input(optionLine)
    while option not in numbers[1:4]:
        print(lines)
        input("Select a correct option... ")
        clear()
        print(menu)
        option = input(optionLine)
    option = int(option)
    if option == 1:
        clear()
        if readText(0) == 'firstTimeOption1 = True\n':
            print(lines)
            print("Put your Excels in the Excel directory")
            # print("")
            print(lines)
            input("Press enter to Confirm... ")
            writeInText(0, "firstTimeOption1 = False\n")
        elif readText(0) == 'firstTimeOption1 = False\n':
            pass
        getDataFrom()
        input("New Data Uploaded... ")
        clear()

    elif option == 2:

        while outOption1 != 1:
            outOption2 = 0
            outOption3 = 0
            outOption4 = 0
            data_Files = []
            error = False
            clear()

            if login == False:
                clear()
                mail, password = enter()
                login = True
                print(lines)
                input("Press enter to continue... ")
                clear()

            print(lines)
            print("From what Excel Data base you want to import data")
            print(lines)

            dataOptions1 = printArray(dataBase())
            print(goBackOption)

            print(lines)
            dataBaseSelection = input(optionLine)

            while dataBaseSelection not in numbers[:dataOptions1+1]:
                print(lines)
                input("Select a correct option... ")
                clear()
                print(lines)
                print("From what Excel Data base you want to import data")
                print(lines)
                dataOptions1 = printArray(dataBase())
                print(goBackOption)
                print(lines)
                dataBaseSelection = input(optionLine)

            dataBaseSelection = int(dataBaseSelection)

            if dataBaseSelection == 0:
                outOption1 = 1
            else:
                selectedDataBase = dataBase()[dataBaseSelection-1]

                dataSelectedFolder = os.path.abspath(
                    os.path.join(os.path.dirname(
                        __file__), 'Data', selectedDataBase))

                mailFolder = os.path.abspath(os.path.join(os.path.dirname(
                    __file__), 'Mails'))

                data_Files = getFilesInFolder(dataSelectedFolder)

                while outOption2 != 1:
                    outOption3 = 0
                    outOption4 = 0
                    clear()
                    print(lines)
                    print("Select what data file have the Mails")
                    print(lines)
                    dataOptions2 = printArray(data_Files)
                    print(goBackOption)
                    print(lines)
                    email_textfile = input(optionLine)

                    while email_textfile not in numbers[:dataOptions2+1]:
                        print(lines)
                        input("Select a correct option... ")
                        clear()
                        print(lines)
                        print("Select what data file have the Mails")
                        print(lines)
                        dataOptions2 = printArray(data_Files)
                        print(goBackOption)
                        print(lines)
                        email_textfile = input(optionLine)

                    email_textfile = int(email_textfile)
                    if email_textfile == 0:
                        outOption2 = 1
                    else:
                        while outOption3 != 1:
                            outOption4 = 0
                            email_database = data_Files[email_textfile-1]
                            for files in data_Files:
                                data_filesRemaining.append(files)
                            data_filesRemaining.remove(email_database)
                            clear()
                            data_FilesWE = []
                            data_FilesWE = removeExtensions(data_filesRemaining)
                            docsInMails = []
                            print(lines)
                            print("Data loaded succesfully, now you can write a Email in word or in a text file with the next variables")
                            print(*data_FilesWE, sep=", ")
                            print(lines)
                            input("Press enter when you got your mail ready in the folder... ")
                            docsInMails = getDocsInFolder(mailFolder)
                            for i in docsInMails:
                                if i[0] == "~" and i[1] == "$":
                                    docsInMails.remove(i)
                            clear()
                            print(lines)
                            print("Select the word document with the mail draft")
                            print(lines)
                            dataOptions3 = printArray(docsInMails)
                            print(goBackOption)
                            print(lines)
                            wordSelection = input(optionLine)

                            while wordSelection not in numbers[:dataOptions3+1]:
                                print(lines)
                                input("Select a correct option... ")
                                clear()
                                print(lines)
                                print("Select the word document with the mail draft")
                                print(lines)
                                dataOptions3 = printArray(docsInMails)
                                print(goBackOption)
                                print(lines)
                                wordSelection = input(optionLine)
                            wordSelection = int(wordSelection)
                            wordSelected = docsInMails[wordSelection-1]
                            if wordSelection == 0:
                                data_filesRemaining = []
                                outOption3 = 1
                            else:
                                while outOption4 != 1:
                                    clear()
                                    print(lines)
                                    print("If you want to send attachments put them in the attachments folder")
                                    print(lines)
                                    input("Press enter to continue... ")
                                    attachmentsFolder = os.path.abspath(os.path.join(os.path.dirname(__file__), 'Attachments'))
                                    attachments = getFilesInFolder(attachmentsFolder)
                                    if len(attachments) == 0:
                                        print(lines)
                                        print("0 attachments founded, sending Mails...")
                                        print(lines)
                                        sendMassiveMails(dataSelectedFolder, data_Files,
                                                    wordSelected, email_database, mail)
                                        outOption4 = 1
                                        outOption3 = 1
                                        outOption2 = 1
                                        outOption1 = 1
                                        break
                                    else:
                                        clear()
                                        print(lines)
                                        print("We found",len(attachments), "How many attachments you will send?")
                                        print("Select 0 if you don't want to send an attachment")
                                        print(lines)
                                        howManyAtt = input("Number: ")
                                        while howManyAtt not in numbers[:len(attachments)+1]:
                                            print(lines)
                                            input("Select a correct option... ")
                                            clear()
                                            print(lines)
                                            print("How many attachments you will send?")
                                            print("Select 0 if you don't want to send an attachment")
                                            print(lines)
                                            howManyAtt = input("Number: ")
                                        howManyAtt = int(howManyAtt)
                                        if howManyAtt == 0:
                                            sendMassiveMails(dataSelectedFolder, data_Files,
                                                        wordSelected, email_database, mail, attachmentsSelected)
                                        else:
                                            for files in range(howManyAtt):
                                                clear()
                                                print(lines)
                                                print("Select the attachments")
                                                print(lines)
                                                dataOption4 = printArray(attachments)
                                                print(goBackOption)
                                                print(lines)
                                                attachmentSelected = input(optionLine)

                                                while attachmentSelected not in numbers[:dataOption4+1]:
                                                    print(lines)
                                                    input("Select a correct option... ")
                                                    clear()
                                                    print(lines)
                                                    print("Select the attachments")
                                                    print(lines)
                                                    dataOption4 = printArray(attachments)
                                                    print(goBackOption)
                                                    print(lines)
                                                    attachmentSelected = input(optionLine)

                                                attachmentSelected = int(attachmentSelected)
                                                if attachmentSelected == 0:
                                                    sendMassiveMails(dataSelectedFolder, data_Files,
                                                        wordSelected, email_database, mail, attachmentsSelected)
                                                else:
                                                    attachmentsSelected.append(attachments[attachmentSelected-1])
                                                    attachments.remove(attachments[attachmentSelected-1])
                                            sendMassiveMails(dataSelectedFolder, data_Files,
                                                            wordSelected, email_database, mail, attachmentsSelected)
                                        outOption4 = 1
                                        outOption3 = 1
                                        outOption2 = 1
                                        outOption1 = 1

    elif option == 3:
        while outOption1 != 1:
            outOption2 = 0
            clear()
            print(loginMenu)
            optionLogin = input(optionLine)
            while optionLogin not in numbers[0:3]:
                print(lines)
                input("Select a correct option... ")
                clear()
                print(loginMenu)
                optionLogin = input(optionLine)
            optionLogin = int(optionLogin)

            if optionLogin ==0:
                outOption1 = 1
            if optionLogin == 1:
                if readText(1) == "AutoSave = [Disabled]\n":
                    savePassword = True
                    writeInText(1, "AutoSave = [Enabled]\n")
                    autosave = readText(1)[:len(readText(1))-1]
                    if login == True:
                        # Encrypting passwords and mails
                        crypter = Fernet(readText(-1).encode())
                        password_encrypted = crypter.encrypt(password.encode())
                        mail_encrypted = crypter.encrypt(mail.encode())
                        writeInText(3, mail_encrypted.decode("utf-8")+"\n")
                        writeInText(4, password_encrypted.decode("utf-8")+"\n")
                    else:
                        mail, password = enter()
                        crypter = Fernet(readText(-1).encode())
                        password_encrypted = crypter.encrypt(password.encode())
                        mail_encrypted = crypter.encrypt(mail.encode())
                        writeInText(3, mail_encrypted.decode("utf-8")+"\n")
                        writeInText(4, password_encrypted.decode("utf-8")+"\n")
                        login = True
                    input("Autosave Enabled... ")
                    outOption1 = 1

                elif readText(1) == "AutoSave = [Enabled]\n":
                    writeInText(1, "AutoSave = [Disabled]\n")
                    writeInText(3, "Mail: \n")
                    writeInText(4, "Password: \n")
                    autosave = readText(1)[:len(readText(1))-1]
                    input("Autosave Disabled... ")
                    outOption1 = 1

            elif optionLogin == 2:
                while outOption2 != 1:
                    clear()
                    htmlFiles = getHTMLSInFolder(htmlsFolder)
                    if len(htmlFiles) == 0:
                        print(lines)
                        print("Please put the HTML files in the folder")
                        print(lines)
                        input("Press enter to continue... ")
                    else:
                        print(lines)
                        print("Select the template you want to use")
                        print(lines)
                        printArray(htmlFiles)
                        print(goBackOption)
                        print(lines)
                        htmlTemplateSelected = input(optionLine)
                        while htmlTemplateSelected not in numbers[0:len(htmlFiles)+1]:
                            clear()
                            print(lines)
                            input("Select a correct option... ")
                            print(lines)
                            printArray(htmlFiles)
                            print(goBackOption)
                            print(lines)
                            htmlTemplateSelected = input(optionLine)
                        htmlTemplateSelected = int(htmlTemplateSelected)
                        if htmlTemplateSelected == 0:
                            outOption2 = 1
                        else:
                            while outOption2 != 1:
                                htmlSelected = htmlFiles[htmlTemplateSelected-1]
                                writeInText(5, htmlSelected+"\n")
                                outOption1 = 1
                                outOption2 = 1
                            outOption1 = 1
                            outOption2 = 1
